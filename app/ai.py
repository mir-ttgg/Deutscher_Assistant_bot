from os import getenv
from pathlib import Path

import docx
from openai import OpenAI

from promt import system_prompt

# def convert_docx_to_txt(docx_path: str) -> str:
#     '''Конвертирует DOCX в текст'''
#     doc = docx.Document(docx_path)
#     text = []
#     for paragraph in doc.paragraphs:
#         if paragraph.text.strip():
#             text.append(paragraph.text)
#     return '\n'.join(text)


# def save_knowledge_base(text: str, output_path: str):
#     '''Сохраняет текст в файл'''
#     with open(output_path, 'w', encoding='utf-8') as f:
#         f.write(text)


# client = OpenAI(api_key=getenv('API_KEY'))


# def upload_knowledge_files(file_paths: List[str]) -> List[str]:
#     '''Загружает файлы в OpenAI и возвращает их ID'''
#     file_ids = []

#     for file_path in file_paths:
#         with open(file_path, 'rb') as f:
#             file = client.files.create(
#                 file=f,
#                 purpose='assistants'
#             )
#             file_ids.append(file.id)

#     return file_ids


# def create_vector_store(name: str, file_ids: List[str]) -> str:
#     '''Создает векторное хранилище с файлами'''
#     vector_store = client.beta.vector_stores.create(
#         name=name,
#         file_ids=file_ids
#     )
#     return vector_store.id


# def create_assistant(vector_store_id: str) -> str:
#     '''Создает ассистента с базой знаний'''

#     # Критически важный промпт
#     system_prompt = promt

#     assistant = client.beta.assistants.create(
#         name='Консультант',
#         instructions=system_prompt,
#         model='gpt-5-mini',
#         tools=[{'type': 'file_search'}],
#         tool_resources={
#             'file_search': {
#                 'vector_store_ids': [vector_store_id]
#             }
#         }
#     )

#     return assistant.id


# docx_files = ['knowledge/document1.docx']
# txt_files = []
# for docx_file in docx_files:
#     text = convert_docx_to_txt(docx_file)
#     txt_path = docx_file.replace('.docx', '.txt')
#     save_knowledge_base(text, txt_path)
#     txt_files.append(txt_path)

# file_ids = upload_knowledge_files(txt_files)
# vector_store_id = create_vector_store("My Knowledge Base", file_ids)
# ASSISTANT_ID = create_assistant(vector_store_id)


# async def get_ai_response(user_message: str, thread_id: str = None) -> tuple[str, str]:
#     """Получает ответ от AI ассистента"""

#     # Создаем или используем существующий thread
#     if not thread_id:
#         thread = client.beta.threads.create()
#         thread_id = thread.id

#     # Добавляем сообщение пользователя
#     client.beta.threads.messages.create(
#         thread_id=thread_id,
#         role="user",
#         content=user_message
#     )

#     # Запускаем ассистента
#     run = client.beta.threads.runs.create(
#         thread_id=thread_id,
#         assistant_id=ASSISTANT_ID
#     )

#     # Ожидаем завершения
#     while run.status in ['queued', 'in_progress']:
#         await asyncio.sleep(0.5)
#         run = client.beta.threads.runs.retrieve(
#             thread_id=thread_id,
#             run_id=run.id
#         )

#     # Получаем ответ
#     messages = client.beta.threads.messages.list(thread_id=thread_id)
#     assistant_message = messages.data[0].content[0].text.value

#     return assistant_message, thread_id


# async def update_knowledge_base(file_path: str):
#     """Обновляет базу знаний новым файлом"""
#     # Конвертируем docx в txt
#     text = convert_docx_to_txt(file_path)
#     txt_path = file_path.replace('.docx', '.txt')
#     save_knowledge_base(text, txt_path)

#     # Загружаем новый файл
#     file_ids = upload_knowledge_files([txt_path])

#     # Создаем новый vector store
#     vector_store_id = create_vector_store("Updated Knowledge Base", file_ids)

#     # Обновляем ассистента
#     client.beta.assistants.update(
#         assistant_id=ASSISTANT_ID,
#         tool_resources={
#             'file_search': {
#                 'vector_store_ids': [vector_store_id]
#             }
#         }
#     )

#     return True


client = OpenAI(api_key=getenv("API_KEY"))

KNOWLEDGE_PATH = Path("knowledge/document1.txt")


def convert_docx_to_txt(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
    return "\n".join(text)


def save_knowledge_base(text: str):
    KNOWLEDGE_PATH.parent.mkdir(exist_ok=True)
    with open(KNOWLEDGE_PATH, "w", encoding="utf-8") as f:
        f.write(text)


def load_knowledge_base() -> str:
    if KNOWLEDGE_PATH.exists():
        with open(KNOWLEDGE_PATH, "r", encoding="utf-8") as f:
            return f.read()
    return ""


async def get_ai_response(user_message: str) -> str:
    """Ответ через Responses API"""

    knowledge_text = load_knowledge_base()

    full_prompt = f"""
{system_prompt}

БАЗА ЗНАНИЙ:
{knowledge_text}

ВОПРОС ПОЛЬЗОВАТЕЛЯ:
{user_message}
"""

    response = client.responses.create(
        model="gpt-5-mini",
        input=full_prompt
    )

    return response.output_text


async def update_knowledge_base(file_path: str):
    """Обновление БЗ"""
    text = convert_docx_to_txt(file_path)
    save_knowledge_base(text)
    return True
